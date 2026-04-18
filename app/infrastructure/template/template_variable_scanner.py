from __future__ import annotations

import re
import zipfile
import xml.etree.ElementTree as ET
from collections import OrderedDict
from dataclasses import dataclass
from typing import Iterable, List

from docx import Document

from app.domain.errors import SourceReadError


PLACEHOLDER_PATTERN = re.compile(r"\{\{\s*([^{}\r\n]+?)\s*\}\}")

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_W_T = f"{{{_W_NS}}}t"


def _normalize_placeholder_braces(s: str) -> str:
    """将全角花括号等替换为半角，便于匹配 {{ }}。"""
    return (
        s.replace("\uff5b", "{")
        .replace("\uff5d", "}")
        .replace("｛", "{")
        .replace("｝", "}")
    )


def _joined_text_from_docx_xml(blob: bytes) -> str:
    """按 XML 文档顺序拼接所有 w:t，修复跨 run 断开的 {{var}}。"""
    try:
        root = ET.fromstring(blob)
    except ET.ParseError:
        return ""
    parts: List[str] = []
    for elem in root.iter(_W_T):
        if elem.text:
            parts.append(elem.text)
    return "".join(parts)


def _ordered_xml_part_names(member_names: List[str]) -> List[str]:
    out: List[str] = []
    if "word/document.xml" in member_names:
        out.append("word/document.xml")
    out.extend(sorted(n for n in member_names if re.fullmatch(r"word/header\d+\.xml", n)))
    out.extend(sorted(n for n in member_names if re.fullmatch(r"word/footer\d+\.xml", n)))
    for n in ("word/footnotes.xml", "word/endnotes.xml"):
        if n in member_names:
            out.append(n)
    return out


@dataclass(frozen=True)
class TemplateScanOptions:
    keep_order: bool = True
    include_header_footer: bool = True


class TemplateVariableScanner:
    def scan(self, template_path: str, options: TemplateScanOptions | None = None) -> List[str]:
        """
        扫描 docx 模板中的 {{变量名}}。
        输出去重后的变量列表（变量名不带 {{ }} 包裹符）。

        优先从 OOXML（document/header/footer 等 XML）内按顺序拼接 w:t 再匹配，
        覆盖：正文、表格、文本框等内嵌 w:t，以及占位符被 Word 拆成多个 run 的情况。
        再用语义层（python-docx）补漏一次。
        """
        options = options or TemplateScanOptions()
        ordered: "OrderedDict[str, None]" = OrderedDict()

        try:
            with zipfile.ZipFile(template_path, "r") as zf:
                names = zf.namelist()
                for part in _ordered_xml_part_names(names):
                    if part not in names:
                        continue
                    blob = zf.read(part)
                    merged = _normalize_placeholder_braces(_joined_text_from_docx_xml(blob))
                    for match in PLACEHOLDER_PATTERN.findall(merged):
                        key = match.strip()
                        if key and key not in ordered:
                            ordered[key] = None
        except zipfile.BadZipFile as e:
            raise SourceReadError(f"无法读取模板 DOCX（非合法 zip）：{template_path}") from e
        except Exception as e:  # noqa: BLE001
            raise SourceReadError(f"无法读取模板 DOCX：{template_path}") from e

        try:
            doc = Document(template_path)
        except Exception as e:  # noqa: BLE001
            raise SourceReadError(f"无法读取模板 DOCX：{template_path}") from e

        for block in self._collect_text_blocks(
            doc,
            include_header_footer=options.include_header_footer,
        ):
            block = _normalize_placeholder_braces(block)
            for match in PLACEHOLDER_PATTERN.findall(block):
                key = match.strip()
                if key and key not in ordered:
                    ordered[key] = None

        if not options.keep_order:
            return sorted(ordered.keys())
        return list(ordered.keys())

    def _collect_text_blocks(self, doc: Document, include_header_footer: bool) -> Iterable[str]:
        for p in doc.paragraphs:
            if p.text:
                yield p.text

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        yield cell.text

        if not include_header_footer:
            return

        for section in doc.sections:
            for p in section.header.paragraphs:
                if p.text:
                    yield p.text
            for t in section.header.tables:
                for row in t.rows:
                    for cell in row.cells:
                        if cell.text:
                            yield cell.text

            for p in section.footer.paragraphs:
                if p.text:
                    yield p.text
            for t in section.footer.tables:
                for row in t.rows:
                    for cell in row.cells:
                        if cell.text:
                            yield cell.text
